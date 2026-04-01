"""
Bullet-line JSON'ından DOCX üreten renderer.

Tüm font, renk, sembol, indent parametreleri sabittir.
LLM veya graph kullanmaz.
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ── SABİT PARAMETRELER ──────────────────────────────────
FONT = "Trebuchet MS"

LOGO_PATH = Path(__file__).parent / "logo.png"
LOGO_WIDTH_CM = 5.27
LOGO_HEIGHT_CM = 1.06
LOGO_ALIGN = "right"

LEVELS = {
    0: {"sym": "❖", "size": 11, "bold": True,  "italic": True},
    1: {"sym": "•", "size": 10, "bold": False, "italic": False},
    2: {"sym": "▪", "size": 10, "bold": False, "italic": False},
    3: {"sym": "–", "size": 10, "bold": False, "italic": False},
}

PLACEHOLDER_TEXT = "(Bir bilgi girilmemiştir.)"
PLACEHOLDER_GRAY = RGBColor(150, 150, 150)

FIXED_BLOCK_COLOR = RGBColor(192, 89, 17)
FIXED_PROJECT_NAME = "Project Name – H_Project"
FIXED_BULLET0 = f"[{FIXED_PROJECT_NAME}]"
FIXED_BULLET1 = "Gelişmeler"
FIXED_ITEMS = ["Bullet 1", "Bullet 2 (gerekirse)", "Bullet 3 (gerekirse)"]

TITLE_TEXT = "HAFTALIK GENEL RAPOR"
TITLE_SIZE = 18
TITLE_COLOR = RGBColor(0, 102, 204)

INSTRUCTION_TEXT = (
    'Bilgi girişi olduğu takdirde "(Bir bilgi girilmemiştir.)" ifadesi silinmelidir.'
)
INSTRUCTION_SIZE = 10
INSTRUCTION_COLOR = RGBColor(255, 0, 0)

INDENT_STEP_IN = 0.45
HANGING_IN = 0.22


# ── YARDIMCI FONKSİYONLAR ───────────────────────────────
def _normalize_to_list(value: Any) -> List[str]:
    if value is None:
        return []
    if isinstance(value, str):
        value = value.strip()
        return [value] if value else []
    if isinstance(value, list):
        out = []
        for item in value:
            if item is None:
                continue
            s = str(item).strip()
            if s:
                out.append(s)
        return out
    s = str(value).strip()
    return [s] if s else []


def _set_run_font(run, *, size: int, bold: bool, italic: bool,
                  color: Optional[RGBColor] = None) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


# ── HEADER LOGO ──────────────────────────────────────────
def _add_logo(doc: Document) -> None:
    if not LOGO_PATH.exists():
        return

    section = doc.sections[0]
    header = section.header

    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.clear()

    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    p.alignment = align_map.get(LOGO_ALIGN, WD_ALIGN_PARAGRAPH.LEFT)

    run = p.add_run()
    run.add_picture(str(LOGO_PATH), width=Cm(LOGO_WIDTH_CM), height=Cm(LOGO_HEIGHT_CM))


# ── BAŞLIK + TALİMAT ────────────────────────────────────
def _add_title_and_instruction(doc: Document) -> None:
    # Başlık
    p1 = doc.add_paragraph()
    r1 = p1.add_run(TITLE_TEXT)
    _set_run_font(r1, size=TITLE_SIZE, bold=True, italic=False, color=TITLE_COLOR)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Talimat
    p2 = doc.add_paragraph()
    r2 = p2.add_run(INSTRUCTION_TEXT)
    _set_run_font(r2, size=INSTRUCTION_SIZE, bold=False, italic=False, color=INSTRUCTION_COLOR)
    r2.underline = True
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()


# ── BULLET EKLEME ────────────────────────────────────────
def _add_bullet(
    doc: Document,
    level: int,
    text: str,
    *,
    override_color: Optional[RGBColor] = None,
    placeholder: bool = False,
) -> None:
    cfg = LEVELS[level]
    p = doc.add_paragraph()

    p.paragraph_format.left_indent = Inches(INDENT_STEP_IN * level)
    p.paragraph_format.first_line_indent = -Inches(HANGING_IN)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    r = p.add_run(f"{cfg['sym']}  {text}")
    _set_run_font(r, size=cfg["size"], bold=cfg["bold"], italic=cfg["italic"])

    if placeholder:
        r.italic = True
        r.font.color.rgb = PLACEHOLDER_GRAY

    if override_color is not None:
        r.font.color.rgb = override_color


# ── SABİT BLOK ───────────────────────────────────────────
def _render_fixed_block(doc: Document) -> None:
    _add_bullet(doc, 0, FIXED_BULLET0, override_color=FIXED_BLOCK_COLOR)
    _add_bullet(doc, 1, FIXED_BULLET1, override_color=FIXED_BLOCK_COLOR)
    _add_bullet(doc, 2, FIXED_ITEMS[0], override_color=FIXED_BLOCK_COLOR)
    for item in FIXED_ITEMS[1:]:
        _add_bullet(doc, 3, item, override_color=FIXED_BLOCK_COLOR)
    doc.add_paragraph()


# ── BULLET SATIRLARI ────────────────────────────────────
def _render_lines(doc: Document, lines: List[Dict[str, Any]]) -> None:
    last_b0 = None

    for ln in lines:
        b0_raw = ln.get("bullet0")
        b0 = b0_raw.strip() if isinstance(b0_raw, str) else None

        bullet1_list = _normalize_to_list(ln.get("bullet1"))
        bullet2_list = _normalize_to_list(ln.get("bullet2"))
        bullet3_list = _normalize_to_list(ln.get("bullet3"))

        if b0 and b0 != last_b0:
            _add_bullet(doc, 0, b0)
            last_b0 = b0

        for b1 in bullet1_list:
            _add_bullet(doc, 1, b1, placeholder=(b1 == PLACEHOLDER_TEXT))

        for b2 in bullet2_list:
            _add_bullet(doc, 2, b2, placeholder=(b2 == PLACEHOLDER_TEXT))

        for b3 in bullet3_list:
            _add_bullet(doc, 3, b3, placeholder=(b3 == PLACEHOLDER_TEXT))


# ══════════════════════════════════════════════════════════
# PUBLIC API
# ══════════════════════════════════════════════════════════

def render_report_to_bytes(bullet_lines: List[Dict[str, Any]]) -> bytes:
    """
    Bullet-line listesini alır, DOCX dosyasını bellekte üretir,
    bytes olarak döndürür.
    """
    doc = Document()

    # 1) Header logo
    _add_logo(doc)

    # 2) Başlık + talimat
    _add_title_and_instruction(doc)

    # 3) Sabit blok
    _render_fixed_block(doc)

    # 4) Gelen bullet satırları (sabit blok varsa çıkar)
    filtered = [
        ln for ln in bullet_lines
        if (ln.get("bullet0") or "").strip() != FIXED_BULLET0
    ]
    _render_lines(doc, filtered)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()